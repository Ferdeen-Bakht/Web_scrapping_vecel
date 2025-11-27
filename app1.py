# app1.py  — Vercel-safe, improved error handling, /tmp file writes
from flask import Flask, render_template, request, send_file
import requests
from lxml import html
from urllib.parse import urljoin, urldefrag
from docx import Document
from docx.shared import Pt, Inches
import os
import traceback

app = Flask(__name__)

# store last scraped content for preview (ephemeral - serverless processes may reset)
last_scraped_content = {}

# Use a single consistent output path in /tmp (suitable for Vercel)
OUTPUT_DOCX_PATH = "/tmp/scraped_content.docx"


def safe_fromstring(content):
    """Parse HTML bytes into an lxml tree with a safe wrapper."""
    try:
        return html.fromstring(content)
    except Exception:
        # try a more forgiving parser fallback
        try:
            parser = html.HTMLParser(recover=True)
            return html.fromstring(content, parser=parser)
        except Exception:
            raise


def scrape_website(base_url):
    """
    Scrape navigation links from the base_url, fetch pages, clean and
    create a docx file containing the successful pages.
    Returns: (output_path_or_None, stats_or_error_dict)
    """
    global last_scraped_content
    stats = {
        "total": 0,
        "success": 0,
        "failed": 0,
        "success_links": [],
        "failed_links": []
    }

    try:
        response = requests.get(base_url, timeout=10)
        response.raise_for_status()
    except Exception as e:
        return None, {"error": f"❌ Failed to fetch base page: {e}"}

    try:
        tree = safe_fromstring(response.content)
    except Exception as e:
        return None, {"error": f"❌ Failed to parse base page: {e}"}

    # Find navigation-like containers
    nav_links = []
    try:
        navbars = tree.xpath('//nav | //ul[contains(@class,"nav")] | //div[contains(@class,"nav")]')
    except Exception:
        navbars = []

    for nav in navbars:
        try:
            links = nav.xpath('.//a/@href')
        except Exception:
            links = []
        for href in links:
            if not href or href.startswith("#"):
                continue
            full_url = urljoin(base_url, href)
            full_url, _ = urldefrag(full_url)
            # Only include same-domain links (basic filter)
            try:
                if full_url not in nav_links and base_url in full_url:
                    nav_links.append(full_url)
            except Exception:
                # fallback: append if unique
                if full_url not in nav_links:
                    nav_links.append(full_url)

    stats["total"] = len(nav_links)
    if not nav_links:
        return None, {"error": "⚠️ No navigation links found."}

    all_content = {}

    for link in nav_links:
        try:
            res = requests.get(link, timeout=10)
            res.raise_for_status()
            try:
                page_tree = safe_fromstring(res.content)
            except Exception:
                # If parsing fails, count as failed and continue
                stats["failed"] += 1
                stats["failed_links"].append(link)
                continue

            # Remove some undesirable tags safely
            try:
                for bad_tag in page_tree.xpath('//script | //style | //noscript | //meta | //footer | //header'):
                    parent = bad_tag.getparent()
                    if parent is not None:
                        parent.remove(bad_tag)
            except Exception:
                # if removal fails, ignore and continue
                pass

            # Get title
            try:
                title = page_tree.xpath('//title/text()')
                title = title[0].strip() if title else "No title"
            except Exception:
                title = "No title"

            # Extract visible text in a conservative way
            visible_text = []
            try:
                for element in page_tree.xpath('//body//*[not(self::script or self::style)]/text()'):
                    if element is None:
                        continue
                    text = element.strip()
                    if text and not text.startswith('{') and not text.startswith('var'):
                        visible_text.append(text)
            except Exception:
                # fallback: try whole-body text
                try:
                    body = page_tree.xpath('//body')[0]
                    text = body.text_content()
                    visible_text = [text.strip()] if text and text.strip() else []
                except Exception:
                    visible_text = []

            clean_text = ' '.join(visible_text)
            clean_text = ' '.join(clean_text.split())

            if clean_text:
                all_content[link] = {"title": title, "text": clean_text}
                stats["success"] += 1
                stats["success_links"].append(link)
            else:
                stats["failed"] += 1
                stats["failed_links"].append(link)

        except Exception:
            stats["failed"] += 1
            stats["failed_links"].append(link)
            # continue with other links

    # Save only successful pages for preview
    last_scraped_content = {k: v for k, v in all_content.items() if v['text'].strip()}

    # Build DOCX only from successful pages
    try:
        doc = Document()
        # Attempt to set page layout; guard against missing internals
        try:
            section = doc.sections[0]
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            # Some python-docx builds don't expose the expected XML nodes; guard that
            try:
                cols = section._sectPr.xpath('./w:cols')
                if cols:
                    cols[0].set('num', '3')
            except Exception:
                pass
        except Exception:
            pass

        # Document title
        title_para = doc.add_paragraph("Scraped Website Content")
        title_para.style = doc.styles.get('Normal', title_para.style)
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(0)
        title_para.paragraph_format.line_spacing = 1
        for run in title_para.runs:
            run.font.size = Pt(10)
            run.bold = True

        for _, data in last_scraped_content.items():
            p_title = doc.add_paragraph(data['title'])
            p_title.style = doc.styles.get('Normal', p_title.style)
            p_title.paragraph_format.space_before = Pt(0)
            p_title.paragraph_format.space_after = Pt(0)
            p_title.paragraph_format.line_spacing = 1
            for run in p_title.runs:
                run.font.size = Pt(10)
                run.bold = True

            p_body = doc.add_paragraph(data['text'])
            p_body.style = doc.styles.get('Normal', p_body.style)
            p_body.paragraph_format.space_before = Pt(0)
            p_body.paragraph_format.space_after = Pt(0)
            p_body.paragraph_format.line_spacing = 1
            for run in p_body.runs:
                run.font.size = Pt(10)

        # Save to /tmp so Vercel allows writing
        output_path = OUTPUT_DOCX_PATH
        doc.save(output_path)
    except Exception as e:
        # If DOCX creation fails, include the traceback in stats for debugging
        tb = traceback.format_exc()
        return None, {"error": f"❌ Failed to create DOCX: {e}", "trace": tb}

    return output_path, stats


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        url = request.form.get("url")
        if not url:
            return render_template("index.html", message="⚠️ Please enter a valid URL.")
        result, stats = scrape_website(url)
        # result is path or None
        if result and os.path.exists(result):
            return render_template(
                "index.html",
                message="✅ Scraping complete!",
                download=True,
                stats=stats
            )
        else:
            # If stats is dict with "trace" it might be large; we only surface the error key
            if isinstance(stats, dict):
                error_msg = stats.get("error") or "⚠️ Scraping failed."
            else:
                error_msg = "⚠️ Scraping failed."
            return render_template("index.html", message=error_msg, stats=stats)
    return render_template("index.html")


@app.route("/preview")
def preview():
    if not last_scraped_content:
        return render_template("index.html", message="⚠️ Please scrape a website first.")
    return render_template("index.html", preview_content=last_scraped_content, download=True)


@app.route("/download")
def download_file():
    path = OUTPUT_DOCX_PATH
    if os.path.exists(path):
        # send_file will stream the /tmp file
        return send_file(path, as_attachment=True)
    else:
        return "⚠️ File not found. Please scrape a website first."


# Keep the app object named `app` for WSGI compatibility (Vercel uses this).
if __name__ == "__main__":
    # When running locally, allow debug; in production Vercel will not use this block
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
