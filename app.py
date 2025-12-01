from flask import Flask, render_template, request, send_file
import requests
from lxml import html
from urllib.parse import urljoin, urldefrag
from docx import Document
from docx.shared import Pt, Inches
import os

app = Flask(__name__)

last_scraped_content = {}


def scrape_website(base_url):
    global last_scraped_content

    stats = {
        "total": 0,
        "success": 0,
        "failed": 0,
        "success_links": [],
        "failed_links": []
    }

    # Fetch base URL
    try:
        response = requests.get(base_url, timeout=10)
        response.raise_for_status()
    except Exception as e:
        return None, {"error": f"Failed to fetch base page: {e}"}

    try:
        tree = html.fromstring(response.content)
    except:
        return None, {"error": "Failed to parse HTML"}

    nav_links = []
    navbars = tree.xpath('//nav | //ul[contains(@class,"nav")] | //div[contains(@class,"nav")]')

    for nav in navbars:
        links = nav.xpath('.//a/@href')
        for href in links:
            if not href or href.startswith("#"):
                continue
            full_url = urljoin(base_url, href)
            full_url, _ = urldefrag(full_url)
            if full_url not in nav_links and base_url in full_url:
                nav_links.append(full_url)

    stats["total"] = len(nav_links)

    if not nav_links:
        return None, {"error": "No navigation links found"}

    all_content = {}

    # Scrape all nav pages
    for link in nav_links:
        try:
            res = requests.get(link, timeout=10)
            res.raise_for_status()
            page_tree = html.fromstring(res.content)

            # Remove unwanted tags
            for bad in page_tree.xpath('//script | //style | //noscript | //meta | //footer | //header'):
                try:
                    bad.getparent().remove(bad)
                except:
                    pass

            # Title
            title = page_tree.xpath('//title/text()')
            title = title[0].strip() if title else "No title"

            # Extract text
            visible = []
            for el in page_tree.xpath('//body//*[not(self::script or self::style)]/text()'):
                text = el.strip()
                if text and not text.startswith("{") and not text.startswith("var"):
                    visible.append(text)

            clean_text = " ".join(" ".join(visible).split())

            if clean_text:
                all_content[link] = {"title": title, "text": clean_text}
                stats["success"] += 1
                stats["success_links"].append(link)
            else:
                stats["failed"] += 1
                stats["failed_links"].append(link)

        except:
            stats["failed"] += 1
            stats["failed_links"].append(link)

    last_scraped_content = all_content

    # ==========================
    #   CREATE DOCX (No styling)
    # ==========================
    try:
        doc = Document()

        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section._sectPr.xpath('./w:cols')[0].set('num', '3')

        for _, data in all_content.items():

            # ---- Heading (plain, not bold, no spacing) ----
            heading = doc.add_paragraph(data["title"])
            heading.style = doc.styles["Normal"]
            heading.paragraph_format.space_before = Pt(0)
            heading.paragraph_format.space_after = Pt(0)
            heading.paragraph_format.line_spacing = 1

            for run in heading.runs:
                run.font.size = Pt(8)
                run.bold = False

            # ---- Body text ----
            body = doc.add_paragraph(data["text"])
            body.style = doc.styles["Normal"]
            body.paragraph_format.space_before = Pt(0)
            body.paragraph_format.space_after = Pt(0)
            body.paragraph_format.line_spacing = 1

            for run in body.runs:
                run.font.size = Pt(8)
                run.bold = False

        output_path = "/tmp/scraped_content.docx"
        doc.save(output_path)

    except Exception as e:
        return None, {"error": f"Failed to create DOCX: {e}"}

    return output_path, stats


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        url = request.form.get("url")
        if not url:
            return render_template("index.html", message="Please enter a valid URL")

        result, stats = scrape_website(url)

        if result and os.path.exists(result):
            return render_template("index.html", message="Scraping complete!", download=True, stats=stats)

        return render_template("index.html", message=stats.get("error", "Scraping failed"))

    return render_template("index.html")


@app.route("/preview")
def preview():
    if not last_scraped_content:
        return render_template("index.html", message="No content scraped yet")

    return render_template("index.html", preview_content=last_scraped_content, download=True)


@app.route("/download")
def download_file():
    path = "/tmp/scraped_content.docx"
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    return "File not found. Please scrape again."


if __name__ == "__main__":
    app.run()
